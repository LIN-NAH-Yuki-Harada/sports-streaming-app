#!/usr/bin/env python3
"""融資依頼書 — 銀行向け標準ビジネスレター形式"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── フォント・ページ設定 ──
FONT_NAME = 'ＭＳ 明朝'
FONT_SIZE = 10.5

style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = Pt(FONT_SIZE)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

def r(p, text, bold=False, size=FONT_SIZE):
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    return run

def right_line(text, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    r(p, text, bold=bold)
    return p

def center_line(text, bold=False, size=FONT_SIZE, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r(p, text, bold=bold, size=size)
    return p

def body(text, space_after=0, indent=0, first_indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    r(p, text)
    return p

def blank(space=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    return p

def bullet(text, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r(p, f'●　{text}')
    return p


# ══════════════════════════════════════
# 文書本体
# ══════════════════════════════════════

# ── 右上: 会社情報 ──
right_line('LIN-NAH株式会社')
right_line('代表取締役　○○ ○○')
right_line('所在地：○○県○○市○○町○-○-○')
right_line('電話番号：○○○-○○○○-○○○○', space_after=12)

# ── タイトル ──
center_line('新規融資に関するご検討依頼', bold=True, size=14, space_before=6, space_after=12)

# ── 前文 ──
body('貴行ますますご繁栄のこととお慶び申し上げます。　平素は格別のお引き立てをいただき、厚く御礼申し上げます。')
body('早速ではございますが、下記の通り、新規融資のご検討をお願い申し上げます。', space_after=12)

# ── 記 ──
center_line('記', bold=False, size=FONT_SIZE, space_before=6, space_after=12)


# ═══ 1. 資金使途 ═══
body('1．資金使途', space_after=6)

body('当社の新規事業として、ローカルスポーツ向けライブ配信プラットフォーム「LIVE SPOtCH（ライブ スポッチ）」の開発・ローンチ・マーケティングにともなう資金として、金7,000,000円の融資を申し込みます。', first_indent=0.5, space_after=6)

bullet('事業背景：当社はマーケティングデザインカンパニーとして、学校・企業向けのデザイン制作、映像制作、Webサイト制作を行ってまいりました。しかし、生成AIの急速な普及によりデザイン業務の受注が減少しており、新たな収益の柱の確立が急務となっております。既存事業で培った映像制作・Web開発の知見を活かし、成長市場であるスポーツ配信領域に参入いたします。', space_after=3)

bullet('事業内容：小中高の部活動・スポーツ少年団の試合を、保護者がスマートフォン1台でTV中継のようにライブ配信できるWebサービスです。スコアオーバーレイ表示、LINE共有、チーム管理、YouTube自動アーカイブ等の機能を備え、月額300円（配信者プラン）・500円（チームプラン）のサブスクリプションモデルで運営いたします。大手スポーツメディアがカバーしないローカルスポーツ市場（対象：保護者・家族1,000万人以上）を対象とし、国内に直接の競合は存在しません。', space_after=3)

bullet('開発状況：すでに全体の約92%の開発が完了し、本番環境で稼働中です。Stripe決済（月額¥300/¥500）、チーム管理・招待機能、YouTube連携（OAuth認証）、パスワードリセット、お問い合わせフォーム等、主要機能はすべて実装済みです。残りのYouTube録画・アップロード機能等を完了し、2026年6〜7月の正式ローンチを予定しております。', space_after=1)
body('　　　アプリURL: https://sports-streaming-app.vercel.app', indent=1.5, space_after=1)
body('　　　LP URL: https://sports-streaming-app.vercel.app/lp', indent=1.5, space_after=3)

bullet('開発コストについて：現在完成している92%分は、一般的な開発会社に外注した場合、約900万円（税込）の開発費に相当します（別紙「開発費用見積もり」参照）。当社では最新のAI開発支援ツールを積極的に活用し、設計・実装・テストの大部分を自社で対応することで、外注費をかけずに開発を進めてまいりました。開発期間はわずか約2週間であり、AIを駆使した高速開発の実績がございます。また、既存のマーケティングデザイン事業を通じて築いたパートナー企業のネットワークを活かし、インフラ構築やサービス連携においても通常よりも低コストでの開発・運用が可能となっております。事業計画上の開発費150万円は、残り8%の開発に必要な外部API利用料・テスト環境費等であり、コーディングそのものの外注費ではございません。', space_after=3)

bullet('資金内訳：マーケティング費2,500,000円（SNS広告・Google広告・PR・販促物）、開発費1,500,000円（外部API利用料・決済連携費・テスト環境）、運転資金1,500,000円（既存事業との並行運営6ヶ月分）、インフラ費1,000,000円（サーバー・配信基盤1年分）、予備費500,000円', space_after=12)


# ═══ 2. 返済方法・返済原資 ═══
body('2．返済方法・返済原資', space_after=6)

bullet('返済期間：元金均等返済　5年', space_after=3)
bullet('返済額：毎月　約122,694円（年利2.0%想定時）', space_after=3)
bullet('返済原資：新規事業の収支計画に基づき、Year 2（2027年7月〜）に営業黒字化を達成し、年間+657万円の営業利益を見込んでおります。Year 2末には累計キャッシュフローがプラスに転換する計画です。したがって、本件借入の返済原資は十分に確保できるものと考えております（詳細は添付の月次損益計算書をご参照ください）。', space_after=3)
bullet('補足：新規事業が計画を下回った場合でも、既存のマーケティングデザイン事業の収入により返済を継続する体制を維持いたします。', space_after=12)


# ═══ 3. 収支見通し（概要） ═══
body('3．収支見通し（概要）', space_after=6)

body('新規事業の3年間の収支見通しは以下のとおりです。', first_indent=0.5, space_after=6)

bullet('Year 1（2026/7〜2027/6）：登録ユーザー8,000人、有料会員1,000人、売上153万円、営業損益▲350万円（先行投資期間）', space_after=3)
bullet('Year 2（2027/7〜2028/6）：登録ユーザー35,000人、有料会員4,550人、売上1,291万円、営業損益+657万円（黒字転換）', space_after=3)
bullet('Year 3（2028/7〜2029/6）：登録ユーザー100,000人、有料会員13,000人、売上5,144万円、営業損益+3,665万円', space_after=12)


# ═══ 4. 添付書類 ═══
body('4．添付書類', space_after=6)

body('銀行の審査を円滑に進めるため、以下の資料を添付いたします。', first_indent=0.5, space_after=6)

bullet('事業計画書（PowerPoint資料）', space_after=3)
bullet('月次損益計算書・返済スケジュール（Excel資料・3年分）', space_after=3)
bullet('試算表（直近分）', space_after=3)
bullet('資金繰り表（今後1年分）', space_after=3)
bullet('会社登記簿謄本', space_after=3)
bullet('直近○期分の決算書（税務申告書の写し）', space_after=3)
bullet('納税証明書', space_after=3)
bullet('代表者の本人確認書類', space_after=3)
bullet('サービスの画面資料（スクリーンショット）', space_after=3)
bullet('開発費用見積もり（参考：開発会社に外注した場合の想定費用）', space_after=18)


# ── 以上 ──
right_line('以上', space_after=0)


# ── 保存 ──
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "融資依頼書.docx")
doc.save(out)
print(f"保存完了: {out}")
