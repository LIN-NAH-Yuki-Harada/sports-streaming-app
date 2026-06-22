#!/usr/bin/env python3
"""融資依頼書 v3 — 埼玉りそな銀行宛 350 万円（協調融資の片翼）

公庫宛 v3 (create_loan_request.py) と同じ書式・同じ事業内容で、
宛先を埼玉りそな銀行に差し替えた協調融資 700 万円のもう一方。

v2 → v3 変更点 (2026-05-29):
- 提出スケジュール再設計: 5/29 本依頼書作成 → 6/1 税理士確認 → 6 月上旬 銀行ご提出 → 6 月中旬 着金・正式ローンチ + 広告開始
- ローンチ時期 2026/6/1 → 2026/6 中旬
- 5 月完璧化進捗を反映: OAuth 通過 (5/13)・YouTube 埋め込み許可・モバイル遷移体感改善・依存ライブラリ最新化 (5/28-29) を完了として記載
- PR マージ件数を最新化 (5/10 時点 126 件 → 5/29 時点 127 件)

公庫宛との主な差分（v3 でも継承）:
- 宛先: 日本政策金融公庫 → 埼玉りそな銀行
- 前文: 既存取引銀行としての関係性を訴求（LIN-NAH の継続取引実績）
- 協調融資の構成説明: 公庫 350 万との並行として表記
- 出力ファイル名: 融資依頼書_埼玉りそな.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ── フォント・ページ設定 ──
FONT_NAME = 'ＭＳ 明朝'
FONT_SIZE = 10.5

style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = Pt(FONT_SIZE)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

def r(p, text, bold=False, size=FONT_SIZE):
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    return run

def right_line(text, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    r(p, text, bold=bold)
    return p

def center_line(text, bold=False, size=FONT_SIZE, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r(p, text, bold=bold, size=size)
    return p

def body(text, space_after=0, indent=0, first_indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    r(p, text)
    return p

def blank(space=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    return p

def bullet(text, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r(p, f'●　{text}')
    return p


# ══════════════════════════════════════
# 文書本体
# ══════════════════════════════════════

# ── 日付 + 宛先 ──
right_line('2026年5月29日', space_after=6)

body('埼玉りそな銀行　御中', space_after=12)

# ── 右上: 会社情報 ──
right_line('LIN-NAH株式会社')
right_line('代表取締役　原田　佑樹')
right_line('所在地：○○県○○市○○町○-○-○')
right_line('電話番号：○○○-○○○○-○○○○', space_after=12)

# ── タイトル ──
center_line('新規融資（協調融資）に関するご検討依頼', bold=True, size=14, space_before=6, space_after=12)

# ── 前文 ──
body('貴行ますますご繁栄のこととお慶び申し上げます。　平素は格別のお引き立てをいただき、厚く御礼申し上げます。')
body('当社は長年にわたり貴行とお取引をさせていただいており、過去にも貴行より融資をいただきながら事業を継続してまいりました。今般、新規事業「LIVE SPOtCH（ライブ スポッチ）」の本格立ち上げに際し、改めてご支援を賜りたく本書をご提出申し上げます。')
body('本依頼書では、貴行様 350 万円 + 日本政策金融公庫様 350 万円の合計 700 万円の協調融資としてご検討いただきたく、下記の通りお願い申し上げます。', space_after=12)

# ── 記 ──
center_line('記', bold=False, size=FONT_SIZE, space_before=6, space_after=12)


# ═══ 1. 資金使途 ═══
body('1．資金使途', space_after=6)

body('当社の新規事業として、ローカルスポーツ向けライブ配信プラットフォーム「LIVE SPOtCH（ライブ スポッチ）」の運営・マーケティング・運転資金として、貴行様より金 3,500,000 円の融資を申し込みます。', first_indent=0.5, space_after=6)

bullet('協調融資の構成：本件は貴行様 350 万円 + 日本政策金融公庫様 350 万円 = 合計 700 万円の協調融資としてご検討をお願いいたします。日本政策金融公庫へは別途 350 万円の融資依頼を進めており、創業・新規事業支援枠での協調融資としての位置づけです。', space_after=3)

bullet('事業背景：当社はマーケティングデザインカンパニーとして、学校・企業向けのデザイン制作、映像制作、Webサイト制作を行ってまいりました。しかし、生成AIの急速な普及によりデザイン業務の受注が減少し、新たな収益の柱の確立が急務となっております。既存事業で培った映像制作・Web開発の知見と、AI 時代に対応した自社内製化能力を活かし、成長市場であるスポーツ配信領域に参入いたします。', space_after=3)

bullet('事業内容：小中高の部活動・スポーツ少年団の試合を、保護者がスマートフォン1台でTV中継品質のライブ配信ができるWebサービスです。スコアボード常時オーバーレイ表示、LINE共有、チーム管理、YouTube Live 同時配信・自動アーカイブ等の機能を備え、月額300円（配信者プラン：ライブ配信専用）・500円（チームプラン：YouTube連携・チーム管理付き）のサブスクリプションモデルで運営いたします。大手スポーツメディアがカバーしないローカルスポーツ市場（対象：保護者・家族 1,000-1,500 万人）を対象とし、国内に直接の競合は存在しません。', space_after=3)

bullet('開発状況：本番環境で 100% 稼働中であり、2026年4月19日に本番稼働を開始、4月25日に Stripe 本番決済稼働、5月1日に YouTube Live 同時配信機能をリリースしました。直近 50 日（2026/4/9-5/29）で計 127 件のプロダクション PR をマージしております。2026年5月10日のブロックシード大会（バレーボール）では 53 分の本番試合配信を 18 名同時視聴で完走しており、本番運用と並行した即応開発体制が機能しております。2026年5月13日には Google OAuth 検証も正式通過し、「Google 公式検証済 OAuth アプリ」として YouTube 連携機能の信頼性を確立、5月28-29日にはチームプラン視聴の YouTube 埋め込み許可・モバイル遷移体感改善・依存ライブラリ最新化を完了し、本番品質を更に引き上げております。', space_after=1)
body('　　　本番URL: https://live-spotch.com', indent=1.5, space_after=3)

bullet('5月の完璧化進捗：当初の完璧化計画 3 本柱（Google OAuth 検証 / 配信プロトコル TCP 化 / 本番運用上の細部仕上げ）のうち、Google OAuth 検証は 2026年5月13日に正式通過済、本番運用上の重要な細部修正（YouTube 視聴埋め込み許可・モバイル遷移体感改善・依存ライブラリ 13 種一括最新化・「アプリを開く」体感空白解消等）も 2026年5月29日までに完了しております。残る配信プロトコル TCP 化（4G 環境下の高画質安定配信化）はローンチ後の継続改善項目として進めてまいります。**6 月中旬の正式ローンチ時点で「使い物になる」レベルは既に達成済**であり、6/1 税理士確認 → 6 月上旬 銀行ご提出 → 6 月中旬 着金・ローンチ + 広告開始の流れで実行いたします。', space_after=3)

bullet('開発コストについて：現在完成している分は、一般的な開発会社に外注した場合、当初の基本機能ベースで 約 1,200 万円（税込）の開発費に相当します。さらに 5/10 までの追加実装（YouTube Live 連携・配信品質改善・運用堅牢化・SEO・LP 最適化等）を含めると、約 1,500 万円（税込）相当の開発を行ったことになります（別紙「開発費用見積もり v2」参照）。当社では Claude 等の最新の AI 開発支援ツールを積極的に活用し、設計・実装・テスト・PR レビューまで自社で対応することで、コストをかけずに開発を進めてまいりました。事業計画上の開発費 210 万円は、LiveKit Cloud Ship プラン（映像配信基盤）の年額、配信プロトコル TCP 化のための中継サーバー費用、AI 開発支援ツールの利用料が中心であり、コーディング工数の外注費ではございません。', space_after=3)

bullet('資金内訳（700 万円全体）：マーケティング費 2,400,000 円（Meta/Google 広告・PRTIMES・SNS・販促物）、開発費 2,100,000 円（LiveKit Cloud Ship + TCP 化中継サーバー + AI 開発支援）、運転資金 1,700,000 円（既存マーケデザイン事業との並行運営 6ヶ月分）、予備費 500,000 円（OAuth 審査追加対応・TCP 化想定外コスト・法務）、インフラ費 300,000 円（YouTube API・Egress 帯域・Vercel Pro）', space_after=3)

bullet('スケジュール：2026年5月29日 本依頼書作成 → 6月1日 顧問税理士確認・最終化 → 6月上旬 貴行へご提出 → 6月中旬 ご審査・着金（協調融資 700 万円）→ 6月中旬 正式ローンチ + 広告運用開始 → 8月末 KPI ゲート 1（有料 5 名）→ 10月初 KPI ゲート 2（有料 20 名）。6月中旬ローンチ時点で広告運用を開始する計画のため、6月上旬の本申込時点で早期審査をお願い申し上げます。', space_after=12)


# ═══ 2. 返済方法・返済原資 ═══
body('2．返済方法・返済原資', space_after=6)

bullet('返済期間：元利均等返済　5年（60回）', space_after=3)
bullet('返済額（貴行様 350 万円分）：毎月　約 61,347 円（年利 2.0% 想定時）', space_after=3)
bullet('返済額（協調融資 700 万円合計）：毎月　約 122,694 円', space_after=3)
bullet('返済開始：2026年7月（正式ローンチ翌月から開始）', space_after=3)
bullet('返済原資：新規事業の収支計画に基づき、Year 2（2027年6月〜）に営業黒字化を達成し、年間 +657 万円の営業利益を見込んでおります。Year 1 末時点で月 35 万円 MRR に対し、月返済額 12.3 万円（協調融資合計）= 返済安全係数 2.85 倍を確保しております。Year 2 末には累計キャッシュフローがプラスに転換する計画です。本件借入の返済原資は十分に確保できるものと考えております（詳細は添付の月次損益計算書をご参照ください）。', space_after=3)
bullet('補足：新規事業が計画を下回った場合でも、既存のマーケティングデザイン事業の収入により返済を継続する体制を維持いたします。役員報酬の調整余力もございます。当社は貴行との既存取引において、過去の返済実績を継続的に履行してまいりました。', space_after=12)


# ═══ 3. 収支見通し（概要） ═══
body('3．収支見通し（概要）', space_after=6)

body('新規事業の3年間の収支見通しは以下のとおりです（6月中旬ローンチ起点・Year 1 = 2026/6〜2027/5）。', first_indent=0.5, space_after=6)

bullet('Year 1（2026/6〜2027/5）：登録ユーザー 8,000 人、有料会員 1,000 人、売上 153 万円、営業損益 ▲350 万円（先行投資期間）', space_after=3)
bullet('Year 2（2027/6〜2028/5）：登録ユーザー 35,000 人、有料会員 4,550 人、売上 1,291 万円、営業損益 +657 万円（黒字転換）', space_after=3)
bullet('Year 3（2028/6〜2029/5）：登録ユーザー 100,000 人、有料会員 13,000 人、売上 5,144 万円、営業損益 +3,665 万円', space_after=12)


# ═══ 4. 添付書類 ═══
body('4．添付書類', space_after=6)

body('銀行の審査を円滑に進めるため、以下の資料を添付いたします。', first_indent=0.5, space_after=6)

bullet('事業計画書 v3（PowerPoint資料・本日 5/29 改訂版・市場規模推計を含む増強版）', space_after=3)
bullet('提案書 v2（マシモ様フィードバック対応版・LIN-NAH 財務改善計画 + LIVE SPOtCH 収益確実性 + 協調融資提案）', space_after=3)
bullet('月次損益計算書 v2・返済スケジュール（Excel資料・3年分・2026/6 起点）', space_after=3)
bullet('試算表（直近分）', space_after=3)
bullet('資金繰り表（今後1年分）', space_after=3)
bullet('会社登記簿謄本', space_after=3)
bullet('直近○期分の決算書（税務申告書の写し）', space_after=3)
bullet('納税証明書', space_after=3)
bullet('代表者の本人確認書類', space_after=3)
bullet('LIVE SPOtCH 画面資料（スクリーンショット・5/10 試合配信実績含む）', space_after=3)
bullet('開発費用見積もり v2（参考：開発会社に外注した場合の想定費用 — 基本機能 約 1,200 万円 / 5/10 実装増分込み 約 1,500 万円）', space_after=3)
bullet('PR マージ実績一覧（GitHub から自動生成・計 127 件・本日 5/29 時点）', space_after=3)
bullet('Google OAuth 検証承認メール（2026-05-13 受信・Project live-spotch / 3 scope 承認）', space_after=18)


# ── 以上 ──
right_line('以上', space_after=0)


# ── 保存 ──
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "融資依頼書_埼玉りそな.docx")
doc.save(out)
print(f"保存完了: {out}")
print(f"")
print(f"=== 文書概要 (埼玉りそな宛 v3 / 2026-05-29) ===")
print(f"宛先: 埼玉りそな銀行")
print(f"申込金額: 350 万円 (協調融資 700 万円のうち埼玉りそな分)")
print(f"本依頼書作成日: 2026年5月29日")
print(f"税理士確認: 2026年6月1日")
print(f"銀行ご提出予定: 2026年6月上旬")
print(f"着金希望: 2026年6月中旬")
print(f"正式ローンチ: 2026年6月中旬 (着金と同時に広告運用開始)")
print(f"返済開始: 2026年7月 (正式ローンチ翌月)")
print(f"月額返済 (埼玉りそな分): 約 61,347 円")
print(f"月額返済 (協調融資合計): 約 122,694 円")
